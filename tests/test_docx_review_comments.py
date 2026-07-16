from io import BytesIO
import zipfile

import pytest
from docx import Document as DocxDocument
from lxml import etree

from docutranslate.ir.document import Document
from docutranslate.translator.ai_translator.docx_translator import (
    DocxTranslator,
    DocxTranslatorConfig,
)


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
NS = {"w": W_NS, "pr": REL_NS}


def build_source_docx() -> bytes:
    doc = DocxDocument()
    paragraph = doc.add_paragraph("Hello world")
    doc.add_comment(paragraph.runs[0], "Existing", author="Human", initials="H")
    doc.add_table(rows=1, cols=1).cell(0, 0).text = "Payment terms"
    doc.sections[0].header.paragraphs[0].text = "Header text"
    stream = BytesIO()
    doc.save(stream)
    return stream.getvalue()


def translate_with_comments(insert_mode: str) -> bytes:
    translator = DocxTranslator(DocxTranslatorConfig(
        skip_translate=True,
        insert_mode=insert_mode,
        separator="\n",
    ))
    source = Document(content=build_source_docx(), suffix=".docx")
    doc, elements, originals = translator._pre_translate(source)
    translated = [f"译文-{index}" for index in range(len(originals))]
    reviews = {index: f"审校意见-{index}" for index in range(len(originals))}
    return translator._after_translate(doc, elements, translated, originals, reviews)


@pytest.mark.parametrize("insert_mode", ["replace", "append", "prepend"])
def test_review_comments_are_word_comments_and_preserve_existing_comments(insert_mode):
    output = translate_with_comments(insert_mode)
    with zipfile.ZipFile(BytesIO(output)) as package:
        comments = etree.fromstring(package.read("word/comments.xml"))
        comment_nodes = comments.xpath(".//w:comment", namespaces=NS)
        authors = [node.get(f"{{{W_NS}}}author") for node in comment_nodes]
        assert authors.count("Human") == 1
        assert authors.count("AI Review") == 3

        ids = [node.get(f"{{{W_NS}}}id") for node in comment_nodes]
        assert len(ids) == len(set(ids))

        all_story_xml = b"".join(
            package.read(name)
            for name in package.namelist()
            if name == "word/document.xml" or name.startswith("word/header")
        )
        for comment_id in ids:
            assert f'w:id="{comment_id}"'.encode() in all_story_xml

        rels = etree.fromstring(package.read("word/_rels/document.xml.rels"))
        assert rels.xpath(
            ".//pr:Relationship[contains(@Type, '/comments')]", namespaces=NS
        )
        header_rels = etree.fromstring(package.read("word/_rels/header1.xml.rels"))
        assert header_rels.xpath(
            ".//pr:Relationship[contains(@Type, '/comments')]", namespaces=NS
        )
        content_types = package.read("[Content_Types].xml")
        assert b"/word/comments.xml" in content_types


def test_disabled_review_does_not_add_ai_review_comments():
    translator = DocxTranslator(DocxTranslatorConfig(skip_translate=True))
    source = Document(content=build_source_docx(), suffix=".docx")
    doc, elements, originals = translator._pre_translate(source)
    output = translator._after_translate(doc, elements, originals, originals, {})
    with zipfile.ZipFile(BytesIO(output)) as package:
        comments = etree.fromstring(package.read("word/comments.xml"))
        authors = comments.xpath(".//w:comment/@w:author", namespaces=NS)
        assert authors == ["Human"]
